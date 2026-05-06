from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\uryuqq\Desktop\fitness\Fitness_App")
OUTPUT = ROOT / "Otchet_FitnessClub_Gabitov.docx"


def set_run_font(run, size=14, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "4472C4")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "2"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=12)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_paragraph(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True,
                  size=14, bold=False, italic=False, spacing_after=0, spacing_before=0,
                  line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_center(doc, text, size=14, bold=False, spacing_after=0, spacing_before=0):
    return add_paragraph(
        doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False,
        size=size, bold=bold, spacing_after=spacing_after, spacing_before=spacing_before
    )


def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    if level == 1:
        p = add_center(doc, text.upper(), size=14, bold=True, spacing_after=10, spacing_before=6)
    else:
        p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=True,
                          size=14, bold=True, spacing_after=6, spacing_before=8)
    return p


def add_list(doc, items):
    for item in items:
        p = add_paragraph(doc, first_indent=True)
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers, rows, widths=None, header_fill="1F4E79", font_size=12):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.allow_autofit = True
    set_table_borders(table)

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color="FFFFFF")
        if widths:
            cell.width = Cm(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cell.width = Cm(widths[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    add_paragraph(doc, "", first_indent=False, spacing_after=4)
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=9.5, name="Consolas")
    add_paragraph(doc, "", first_indent=False, spacing_after=4)


def add_toc_entry(doc, title, page):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(title)
    set_run_font(run)
    run = p.add_run(f"\t{page}")
    set_run_font(run)


def build_report():
    doc = Document()
    configure_document(doc)

    # Title page
    add_center(doc, "Министерство образования и науки Республики Татарстан", size=14)
    add_center(doc, "ГАПОУ «Казанский радиомеханический колледж»", size=14)
    for _ in range(5):
        add_paragraph(doc, "", first_indent=False)
    add_center(doc, "ОТЧЕТ", size=16, bold=True, spacing_after=4)
    add_center(doc, "ПО ПРАКТИЧЕСКОЙ ПОДГОТОВКЕ", size=16, bold=True, spacing_after=4)
    add_center(doc, "ПМ.01 УП.01.01", size=14, bold=True)
    add_center(doc, "РАЗРАБОТКА МОДУЛЕЙ ПРОГРАММНОГО ОБЕСПЕЧЕНИЯ", size=14, bold=True)
    add_center(doc, "ДЛЯ КОМПЬЮТЕРНЫХ СИСТЕМ", size=14, bold=True)
    add_paragraph(doc, "", first_indent=False)
    add_center(doc, "на тему:", size=14)
    add_center(doc, "Разработка клиент-серверного веб-приложения «PulseBoard»", size=14, bold=True)
    add_center(doc, "для управления фитнес-клубом", size=14, bold=True)

    for _ in range(4):
        add_paragraph(doc, "", first_indent=False)
    p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False, line_spacing=1.2)
    for line in [
        "Выполнил обучающийся",
        "Габитов Амир Ирекович",
        "группы 320",
        "специальность 09.02.07",
        "Информационные системы и программирование",
        "",
        "Руководитель практики",
        "______________________________",
    ]:
        if line:
            run = p.add_run(line)
            set_run_font(run, size=14)
        p.add_run().add_break()
    for _ in range(4):
        add_paragraph(doc, "", first_indent=False)
    add_center(doc, "Казань, 2026", size=14)

    # Contents
    doc.add_page_break()
    add_heading(doc, "СОДЕРЖАНИЕ", level=1)
    toc_entries = [
        ("ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ НА ТЕМУ: веб-приложение «PulseBoard»", "3"),
        ("1 ОБОСНОВАНИЕ ВЫБОРА СРЕДСТВ РАЗРАБОТКИ", "4"),
        ("1.1 Обоснование выбора языка программирования и платформы API", "4"),
        ("1.2 Обоснование выбора клиентской технологии", "5"),
        ("1.3 Обоснование выбора базы данных и ORM", "6"),
        ("1.4 Обоснование выбора средств обмена данными и уведомлений", "7"),
        ("2 РАЗРАБОТКА ПРОГРАММНОГО ПРОДУКТА", "8"),
        ("2.1 Анализ предметной области", "8"),
        ("2.2 Проектирование базы данных", "9"),
        ("2.3 Архитектура и структура проекта", "10"),
        ("2.4 Реализация интерфейсов и функционала", "11"),
        ("2.5 Тестирование программного продукта", "13"),
        ("2.6 Руководство пользователя", "14"),
        ("ЗАКЛЮЧЕНИЕ", "15"),
    ]
    for title, page in toc_entries:
        add_toc_entry(doc, title, page)

    # Assignment
    add_heading(doc, "ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ НА ТЕМУ: веб-приложение «PulseBoard»", level=1, page_break=True)
    add_paragraph(doc, "Темой индивидуального задания является разработка клиент-серверного веб-приложения для управления фитнес-клубом. Программный продукт предназначен для учета клиентов, абонементов, тренировок, посещений, взаимодействия тренеров и клиентов, а также для получения сводной информации по работе клуба.")
    add_paragraph(doc, "Вариант 5 предусматривает реализацию коммуникаций между тренером и клиентом с использованием чата и push-уведомлений. Дополнительно в системе реализованы основные блоки управления фитнес-клубом, так как они необходимы для полноценной проверки работы приложения.")
    add_table(
        doc,
        ["№", "Функциональный блок", "Реализация в проекте"],
        [
            ("1", "Управление абонементами", "Продажа, продление, заморозка, разморозка и просмотр статусов абонементов клиентов."),
            ("2", "Расписание тренировок", "Создание групповых тренировок, запись клиентов, контроль вместимости занятия."),
            ("3", "Индивидуальные планы", "Назначение тренера, план тренировок, план питания и процент прогресса клиента."),
            ("4", "Учет посещений", "Отметка входа и выхода клиента, запрет повторного входа до завершения текущего посещения."),
            ("5", "Коммуникации", "Чат тренер-клиент на SignalR, хранение сообщений в базе данных, уведомления о событиях."),
            ("6", "Отчетность", "Сводные показатели по клиентам, абонементам, посещениям, тренировкам и выручке."),
            ("7", "Администрирование", "Единая панель для ролей администратора, менеджера, тренера и клиента."),
        ],
        widths=[1.1, 5.2, 10],
        font_size=11,
    )

    # Section 1
    add_heading(doc, "1 ОБОСНОВАНИЕ ВЫБОРА СРЕДСТВ РАЗРАБОТКИ", level=1, page_break=True)
    add_heading(doc, "1.1 Обоснование выбора языка программирования и платформы API", level=2)
    add_paragraph(doc, "Для разработки серверной части был выбран язык программирования C# и платформа ASP.NET Core. Данный стек хорошо подходит для построения REST API, так как содержит встроенные средства маршрутизации, внедрения зависимостей, сериализации JSON, подключения Swagger/OpenAPI и работы с базами данных через Entity Framework Core.")
    add_paragraph(doc, "ASP.NET Core позволяет удобно разделить проект на контроллеры, сервисы, модели данных и классы запросов. В разработанном приложении контроллер отвечает за прием HTTP-запроса, сервис содержит основную бизнес-логику, а доступ к данным выполняется через FitnessDbContext. Такой подход соответствует учебному примеру и не усложняет проект отдельным слоем репозиториев.")
    add_paragraph(doc, "API построено на .NET 8.0. Эта версия платформы является актуальной для учебных проектов, поддерживает современный синтаксис C# и содержит все необходимые библиотеки для разработки веб-приложения.")

    add_heading(doc, "1.2 Обоснование выбора клиентской технологии", level=2, page_break=True)
    add_paragraph(doc, "Клиентская часть реализована на Blazor WebAssembly. Выбор Blazor позволяет использовать C# не только на сервере, но и в браузере. За счет этого можно описывать модели запросов и логику обработки данных в привычном для .NET стиле, не переходя полностью на JavaScript.")
    add_paragraph(doc, "Интерфейс приложения выполнен в виде одностраничной панели управления. После авторизации пользователь видит набор блоков, доступных его роли: администратор управляет всеми данными, менеджер работает с клиентами и абонементами, тренер ведет расписание, планы и чат, клиент просматривает свои тренировки, абонемент, посещения и сообщения.")
    add_paragraph(doc, "Обмен данными между клиентом и API выполняется через HttpClient. Для хранения текущей сессии используется browser storage, а сервис AuthService отвечает за вход, выход и проверку состояния пользователя.")

    add_heading(doc, "1.3 Обоснование выбора базы данных и ORM", level=2, page_break=True)
    add_paragraph(doc, "В качестве основной СУБД используется PostgreSQL. Она является бесплатной, надежной и хорошо подходит для хранения структурированных данных: пользователей, ролей, абонементов, тренировок, посещений, сообщений и уведомлений.")
    add_paragraph(doc, "Для работы с базой данных применяется Entity Framework Core и провайдер Npgsql.EntityFrameworkCore.PostgreSQL. ORM позволяет описывать таблицы в виде C#-классов, выполнять миграции, задавать связи между сущностями и работать с данными через DbSet.")
    add_paragraph(doc, "В проекте создан класс FitnessDbContext, в котором объявлены основные наборы данных и связи между таблицами. Для дат в PostgreSQL используется тип timestamp without time zone, что позволяет избежать ошибок при сохранении локального времени посещений, тренировок и сообщений.")

    add_heading(doc, "1.4 Обоснование выбора средств обмена данными и уведомлений", level=2, page_break=True)
    add_paragraph(doc, "REST API используется для операций, где пользователь явно выполняет действие: создать абонемент, продлить срок, записать клиента на тренировку, получить список клиентов или отчет. Формат JSON выбран как стандартный способ обмена данными между сервером и браузером.")
    add_paragraph(doc, "Для коммуникаций в реальном времени используется SignalR. Эта технология позволяет серверу отправлять сообщения клиенту без постоянного ручного обновления страницы. В приложении SignalR применяется для чата тренер-клиент и уведомлений о событиях.")
    add_paragraph(doc, "Документация API подключена через Swagger/OpenAPI. Это упрощает проверку контроллеров во время разработки, так как можно отправлять запросы к серверу прямо из браузера и видеть структуру входных и выходных данных.")
    add_table(
        doc,
        ["Компонент", "Технология", "Назначение"],
        [
            ("Backend API", "C# / ASP.NET Core 8", "REST API, контроллеры, сервисы, обработка бизнес-логики."),
            ("Frontend", "Blazor WebAssembly", "Клиентская панель управления фитнес-клубом."),
            ("Real-time", "SignalR", "Чат и уведомления без обновления страницы."),
            ("База данных", "PostgreSQL", "Хранение учетных данных, расписания, посещений и сообщений."),
            ("ORM", "Entity Framework Core", "Описание моделей, миграции, запросы к базе данных."),
            ("Авторизация", "JWT + сессии", "Вход пользователей и разграничение ролей."),
            ("Документация", "Swagger / OpenAPI", "Проверка и описание HTTP-методов."),
            ("Формат данных", "JSON", "Обмен данными между API и Blazor-клиентом."),
        ],
        widths=[4.2, 4.8, 7.4],
        font_size=11,
    )

    # Section 2
    add_heading(doc, "2 РАЗРАБОТКА ПРОГРАММНОГО ПРОДУКТА", level=1, page_break=True)
    add_heading(doc, "2.1 Анализ предметной области", level=2)
    add_paragraph(doc, "Фитнес-клуб ежедневно работает с несколькими потоками данных: заявки клиентов, продажа и продление абонементов, расписание занятий, посещения, работа тренеров и коммуникации с клиентами. При ручном учете эти процессы становятся разрозненными: менеджер хранит клиентов отдельно, тренер ведет планы отдельно, а клиенту приходится уточнять информацию лично.")
    add_paragraph(doc, "Разрабатываемая система объединяет эти процессы в одну панель. Пользователь входит в систему под своей ролью и получает только те действия, которые соответствуют его обязанностям. Это снижает вероятность ошибок, ускоряет работу сотрудников и делает взаимодействие с клиентом более прозрачным.")
    add_table(
        doc,
        ["Роль", "Основные задачи в системе"],
        [
            ("Администратор", "Просмотр общей панели, управление пользователями, клиентами, тренерами, абонементами, тренировками, отчетами и сообщениями."),
            ("Менеджер по продажам", "Добавление клиентов, продажа и продление абонементов, контроль статуса абонемента, просмотр отчетов по продажам."),
            ("Тренер", "Создание тренировок, запись клиентов, ведение индивидуальных планов, общение с клиентами в чате."),
            ("Клиент", "Просмотр своего абонемента, тренировок, плана, отметка посещения, переписка с тренером."),
        ],
        widths=[4.3, 12.2],
        font_size=11,
    )
    add_paragraph(doc, "Для предметной области были выделены следующие основные сущности: пользователь, роль, клиент, тренер, абонемент, тренировка, индивидуальный план, посещение, сообщение чата, уведомление и сессия. Эти сущности покрывают все функциональные блоки технического задания.")

    add_heading(doc, "2.2 Проектирование базы данных", level=2, page_break=True)
    add_paragraph(doc, "База данных проектировалась с учетом связи между сотрудниками клуба и клиентами. Пользователь хранит учетные данные и роль. Если пользователь является клиентом или тренером, его запись дополнительно связывается с таблицами Clients или Trainers.")
    add_paragraph(doc, "Абонемент связан с клиентом, тренировка связана с тренером, индивидуальный план связан одновременно с клиентом и тренером. Посещение фиксирует время входа и выхода клиента. Сообщение чата хранит отправителя, текст, дату отправки и пару клиент-тренер.")
    add_table(
        doc,
        ["Таблица", "Назначение", "Ключевые поля"],
        [
            ("Roles", "Список ролей пользователей.", "id_Role, Name"),
            ("Users", "Учетные записи для входа в систему.", "id_User, Name, Email, Password, Role_Id, Client_Id, Trainer_Id"),
            ("Sessions", "Активные сессии пользователей.", "id_Session, User_Id, Token, CreatedAt"),
            ("Clients", "Клиенты фитнес-клуба.", "Id, FullName, Phone"),
            ("Trainers", "Тренеры и их специализация.", "Id, FullName, Specialization"),
            ("Memberships", "Абонементы клиентов.", "Id, ClientId, Type, StartDate, EndDate, Status"),
            ("Workouts", "Групповые и персональные занятия.", "Id, Title, TrainerId, StartAt, Capacity, ClientIds"),
            ("Plans", "Индивидуальные планы тренировок и питания.", "Id, ClientId, TrainerId, TrainingPlan, NutritionPlan, ProgressPercent"),
            ("Visits", "Фиксация входа и выхода клиента.", "Id, ClientId, EnteredAt, ExitedAt, AccessType"),
            ("ChatMessages", "История сообщений чата.", "Id, TrainerId, ClientId, SenderRole, Text, SentAt"),
            ("Notifications", "Push-уведомления о событиях.", "Id, ClientId, Text, SentAt, IsRead"),
        ],
        widths=[3.5, 6.5, 6.5],
        font_size=10.5,
    )
    add_paragraph(doc, "Схема базы данных создается миграцией Entity Framework Core. При запуске миграции формируются таблицы, первичные ключи, внешние ключи и начальные данные для демонстрации: роли, несколько клиентов, тренеров, пользователей, абонементов, тренировок и планов.")

    add_heading(doc, "2.3 Архитектура и структура проекта", level=2, page_break=True)
    add_paragraph(doc, "Проект состоит из двух приложений: серверного API и клиентского Blazor WebAssembly. Такой вариант разделения удобен для учебного проекта: сервер можно проверять через Swagger, а клиент работает как отдельная веб-панель и обращается к API по HTTP.")
    add_paragraph(doc, "Серверная часть выполнена в простом стиле: контроллер принимает запрос, вызывает соответствующий сервис, сервис работает с FitnessDbContext и возвращает результат. В проекте используются интерфейсы сервисов, чтобы зависимости подключались через встроенный контейнер ASP.NET Core.")
    add_table(
        doc,
        ["Папка / файл", "Назначение"],
        [
            ("Controllers", "HTTP-контроллеры для клиентов, тренеров, абонементов, тренировок, посещений, планов, отчетов и коммуникаций."),
            ("Services", "Классы бизнес-логики: работа с пользователями, абонементами, тренировками, посещениями, отчетами и сообщениями."),
            ("Interfaces", "Интерфейсы сервисов, которые регистрируются в Program.cs."),
            ("Models", "Модели сущностей базы данных."),
            ("Requests", "DTO-классы для входных данных API."),
            ("Data/FitnessDbContext.cs", "Контекст Entity Framework Core и настройка связей между таблицами."),
            ("Hubs", "SignalR-хабы для чата и уведомлений."),
            ("Migrations", "Миграции базы данных."),
            ("client/Pages", "Страницы Blazor: вход, регистрация и основная панель."),
            ("client/Services", "Сервисы клиента: AuthService, ApiService, SessionStorageService."),
        ],
        widths=[5.4, 11.1],
        font_size=11,
    )
    add_paragraph(doc, "В файле Program.cs API регистрируются контроллеры, Swagger, CORS, SignalR, FitnessDbContext и сервисы предметной области. После сборки приложения маршруты контроллеров и SignalR-хабы становятся доступными клиентской части.")
    add_code_block(doc, """builder.Services.AddDbContext<FitnessDbContext>(options =>
    options.UseNpgsql(builder.Configuration.GetConnectionString("FitnessDbString")));

builder.Services.AddScoped<IUserServices, UserServices>();
builder.Services.AddScoped<IMembershipServices, MembershipServices>();
builder.Services.AddScoped<ICommunicationServices, CommunicationServices>();

app.MapControllers();
app.MapHub<NotificationHub>("/hubs/notifications");
app.MapHub<ChatHub>("/hubs/chat");""")
    add_paragraph(doc, "Клиентское приложение подключает HttpClient с адресом API и собственные сервисы. Это похоже на учебный пример: страница Blazor вызывает сервис, сервис отправляет запрос на сервер, а результат отображается в интерфейсе.")

    add_heading(doc, "2.4 Реализация интерфейсов и функционала", level=2, page_break=False)
    add_paragraph(doc, "Основной экран приложения представляет собой панель управления. Все подписи интерфейса выполнены на русском языке, чтобы пользователь не сталкивался со смешением терминов. После входа на странице отображается приветствие, текущая роль и доступные блоки.")
    add_paragraph(doc, "Авторизация выполняется через форму входа. Пользователь вводит email и пароль, сервер проверяет данные, создает сессию и возвращает токен. Клиент сохраняет сессию в хранилище браузера. При нажатии кнопки «Выйти» данные сессии удаляются, и пользователь возвращается на страницу входа.")
    add_paragraph(doc, "В системе используются следующие демонстрационные роли: Administrator, SalesManager, Trainer и Client. Для каждой роли предусмотрен свой набор действий, но одна страница Home.razor используется как общая панель, что делает проект компактным и понятным.")
    add_table(
        doc,
        ["Функциональность", "Описание реализации"],
        [
            ("Абонементы", "Менеджер или администратор создает абонемент, продлевает срок, замораживает или размораживает его. Продление на 30 дней не меняет статус замороженного абонемента."),
            ("Расписание", "Тренер создает занятие, задает дату, вместимость и список записанных клиентов. Система не допускает запись сверх лимита."),
            ("Планы", "Тренер назначает клиенту индивидуальный план тренировок и питания, а также процент прогресса."),
            ("Посещения", "Клиент отмечает вход и выход. Пока посещение активно, повторная кнопка входа заблокирована."),
            ("Коммуникации", "Сообщения передаются через SignalR и дополнительно сохраняются в таблице ChatMessages."),
            ("Уведомления", "При важных действиях клиент получает уведомление, которое хранится в базе данных."),
            ("Отчеты", "Панель показывает количество клиентов, активных абонементов, тренировок, посещений и финансовые показатели."),
        ],
        widths=[4.5, 12],
        font_size=11,
    )
    add_paragraph(doc, "Для REST API реализованы контроллеры, которые соответствуют функциональным блокам. Ниже приведены основные маршруты, используемые клиентским приложением.")
    add_table(
        doc,
        ["Контроллер", "Примеры методов", "Назначение"],
        [
            ("UserController", "POST api/user/login, POST api/user/register", "Регистрация, вход и получение данных пользователя."),
            ("ClientsController", "GET api/clients, POST api/clients", "Просмотр и создание клиентов."),
            ("TrainersController", "GET api/trainers, POST api/trainers", "Просмотр и создание тренеров."),
            ("MembershipsController", "GET api/memberships, POST, PUT freeze/unfreeze/extend", "Управление абонементами."),
            ("WorkoutsController", "GET api/workouts, POST, PUT enroll", "Расписание тренировок и запись клиентов."),
            ("PlansController", "GET api/plans, POST/PUT", "Индивидуальные планы тренировок и питания."),
            ("VisitsController", "GET api/visits, POST enter, PUT exit", "Учет посещений клиентов."),
            ("CommunicationController", "GET api/communication/chat, POST api/communication/notify", "История сообщений и уведомления."),
            ("ReportsController", "GET api/reports/summary", "Сводная отчетность по клубу."),
        ],
        widths=[4.1, 6.9, 5.5],
        font_size=10,
    )
    add_paragraph(doc, "Чат реализован отдельным SignalR-хабом ChatHub. При подключении пользователь определяется по токену сессии. После этого он добавляется в группу администратора, тренера или клиента. Благодаря группам сообщение сразу получают обе стороны переписки.")
    add_code_block(doc, """public async Task SendChatMessage(SendChatRequest request)
{
    var user = GetCurrentUser();
    var message = new ChatMessage
    {
        TrainerId = trainerId,
        ClientId = clientId,
        SenderRole = senderRole,
        Text = request.Text.Trim(),
        SentAt = DateTime.SpecifyKind(DateTime.Now, DateTimeKind.Unspecified)
    };

    _context.ChatMessages.Add(message);
    await _context.SaveChangesAsync();

    await Clients.Group(TrainerGroup(trainerId)).SendAsync("ReceiveChatMessage", message);
    await Clients.Group(ClientGroup(clientId)).SendAsync("ReceiveChatMessage", message);
}""")
    add_paragraph(doc, "Такой подход позволяет не обновлять страницу вручную. Если тренер отправляет сообщение клиенту, клиент получает его сразу после доставки события ReceiveChatMessage. При повторном открытии панели история сообщений загружается через REST-метод.")

    add_heading(doc, "2.5 Тестирование программного продукта", level=2, page_break=True)
    add_paragraph(doc, "Тестирование выполнялось вручную через интерфейс Blazor, Swagger и прямую проверку данных в PostgreSQL. Основное внимание уделялось авторизации, корректности сохранения данных и соответствию действий пользователя его роли.")
    add_table(
        doc,
        ["№", "Проверяемое действие", "Ожидаемый результат", "Результат"],
        [
            ("1", "Вход под администратором", "Открывается общая панель управления.", "Успешно"),
            ("2", "Вход под клиентом", "Клиент видит свой абонемент, посещения, план и чат.", "Успешно"),
            ("3", "Создание клиента", "Запись появляется в таблице Clients.", "Успешно"),
            ("4", "Продление абонемента на 30 дней", "Дата окончания увеличивается, статус не меняется.", "Успешно"),
            ("5", "Заморозка и разморозка абонемента", "Статус корректно меняется на Frozen и Active.", "Успешно"),
            ("6", "Отметка входа клиента", "Создается активное посещение, кнопка входа становится неактивной.", "Успешно"),
            ("7", "Отметка выхода клиента", "В активном посещении заполняется ExitedAt.", "Успешно"),
            ("8", "Запись клиента на тренировку", "Клиент добавляется в список занятия при наличии мест.", "Успешно"),
            ("9", "Отправка сообщения в чате", "Сообщение появляется у тренера и клиента без обновления страницы.", "Успешно"),
            ("10", "Обновление страницы после авторизации", "Сессия сохраняется, панель загружается повторно.", "Успешно"),
        ],
        widths=[1, 5.2, 7.2, 3],
        font_size=10,
    )
    add_paragraph(doc, "Во время проверки были исправлены ошибки, связанные с бесконечной загрузкой панели, подключением к PostgreSQL, повторной отметкой входа клиента и поведением абонемента при продлении после заморозки. Также был упрощен клиентский код авторизации, чтобы он соответствовал выбранному стилю проекта.")
    add_paragraph(doc, "Сборка API и Blazor-клиента выполнялась средствами dotnet build. Ошибки компиляции после исправлений отсутствуют, что подтверждает корректность подключения сервисов, контроллеров, моделей и клиентских зависимостей.")

    add_heading(doc, "2.6 Руководство пользователя", level=2, page_break=True)
    add_paragraph(doc, "Для начала работы необходимо запустить серверное приложение Fitness_Api и клиентское приложение Fitness_Client. API по умолчанию доступно по адресу http://localhost:5294, а клиент открывается через адрес, который выводит Blazor Dev Server.")
    add_paragraph(doc, "Пользователь открывает страницу входа, вводит email и пароль. После успешной авторизации он попадает на главную панель. Если пользователь нажимает «Выйти», сессия удаляется, и повторный доступ к панели требует нового входа.")
    add_table(
        doc,
        ["Роль", "Порядок работы"],
        [
            ("Администратор", "Войти под учетной записью администратора, проверить сводные показатели, при необходимости открыть списки клиентов, тренеров, абонементов, тренировок, посещений и сообщений."),
            ("Менеджер", "Добавить клиента, оформить абонемент, продлить срок действия или изменить статус. Для проверки результата открыть список абонементов."),
            ("Тренер", "Создать тренировку, записать клиента, заполнить индивидуальный план, открыть чат и отправить сообщение клиенту."),
            ("Клиент", "Проверить данные абонемента, отметить вход при посещении клуба, после завершения нажать «Отметить выход», открыть чат с тренером."),
        ],
        widths=[3.7, 12.8],
        font_size=11,
    )
    add_paragraph(doc, "Если требуется проверить API отдельно от интерфейса, можно открыть Swagger. Через Swagger удобно убедиться, что контроллеры возвращают данные и принимают запросы в формате JSON.")
    add_paragraph(doc, "Для проверки базы данных можно открыть PostgreSQL и посмотреть таблицы Clients, Memberships, Workouts, Visits и ChatMessages. После выполнения действий в интерфейсе соответствующие записи должны появляться или изменяться в этих таблицах.")

    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1, page_break=True)
    add_paragraph(doc, "В ходе выполнения индивидуального задания было разработано клиент-серверное веб-приложение «PulseBoard» для управления фитнес-клубом. Приложение реализует основные процессы предметной области: учет клиентов, абонементов, тренировок, индивидуальных планов, посещений, сообщений и отчетных показателей.")
    add_paragraph(doc, "Серверная часть выполнена на ASP.NET Core и предоставляет REST API. Для хранения данных используется PostgreSQL, а работа с таблицами организована через Entity Framework Core. Клиентская часть разработана на Blazor WebAssembly и взаимодействует с API через HttpClient.")
    add_paragraph(doc, "Особое внимание уделено варианту 5: коммуникациям между тренером и клиентом. Для этого реализован чат на SignalR, который позволяет получать сообщения в реальном времени и сохранять историю переписки в базе данных.")
    add_paragraph(doc, "В результате был получен работоспособный программный продукт, соответствующий технологическому стеку задания и пригодный для демонстрации основных функций управления фитнес-клубом.")

    doc.core_properties.title = "Отчет по веб-приложению PulseBoard"
    doc.core_properties.author = "Габитов Амир Ирекович"
    doc.core_properties.subject = "Платформа управления фитнес-клубом"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(path)
